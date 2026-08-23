"""Export report HTML to PDF / DOCX / HTML."""
from __future__ import annotations

import io
import re
from html.parser import HTMLParser
from typing import Literal

from docx import Document
from docx.shared import Inches, Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import Image as RLImage, Paragraph, SimpleDocTemplate, Spacer

ExportFormat = Literal["pdf", "docx", "html"]


def _strip_tags(html: str) -> str:
    text = re.sub(r"<br\s*/?>", "\n", html, flags=re.I)
    text = re.sub(r"</p>", "\n", text, flags=re.I)
    text = re.sub(r"</h[1-6]>", "\n", text, flags=re.I)
    text = re.sub(r"</li>", "\n", text, flags=re.I)
    text = re.sub(r"<[^>]+>", "", text)
    return re.sub(r"\n{3,}", "\n\n", text).strip()


def wrap_report_html(body_html: str, title: str) -> str:
    return f"""<!DOCTYPE html>
<html lang="uk">
<head>
  <meta charset="utf-8" />
  <title>{title}</title>
  <style>
    body {{ font-family: Helvetica, Arial, sans-serif; font-size: 12pt; color: #1a1a1a; margin: 2cm; line-height: 1.45; }}
    h1 {{ font-size: 20pt; margin-bottom: 0.4em; }}
    h2 {{ font-size: 14pt; margin-top: 1.2em; margin-bottom: 0.4em; }}
    p {{ margin: 0.4em 0; }}
    ul {{ margin: 0.4em 0 0.8em 1.2em; }}
    img {{ max-width: 100%; height: auto; margin: 0.8em 0; }}
    table {{ border-collapse: collapse; width: 100%; margin: 0.8em 0; }}
    th, td {{ border: 1px solid #ccc; padding: 6px 8px; text-align: left; }}
  </style>
</head>
<body>
{body_html}
</body>
</html>"""


def export_html_bytes(body_html: str, title: str) -> bytes:
    return wrap_report_html(body_html, title).encode("utf-8")


def _export_pdf_weasyprint(body_html: str, title: str) -> bytes:
    from weasyprint import HTML

    full = wrap_report_html(body_html, title)
    return HTML(string=full).write_pdf()


class _SimpleHTMLCollector(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.blocks: list[tuple[str, str]] = []
        self._tag_stack: list[str] = []
        self._buf = ""

    def handle_starttag(self, tag: str, attrs) -> None:
        if tag in ("h1", "h2", "h3", "p", "li"):
            self._tag_stack.append(tag)
            self._buf = ""
        elif tag == "br":
            self._buf += "\n"
        elif tag == "img":
            src = dict(attrs).get("src", "")
            if src.startswith("data:image"):
                self.blocks.append(("img", src))

    def handle_endtag(self, tag: str) -> None:
        if tag in ("h1", "h2", "h3", "p", "li") and self._tag_stack and self._tag_stack[-1] == tag:
            self._tag_stack.pop()
            text = self._buf.strip()
            if text:
                self.blocks.append((tag, text))
            self._buf = ""

    def handle_data(self, data: str) -> None:
        if self._tag_stack:
            self._buf += data


def _export_pdf_reportlab(body_html: str, title: str) -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=2 * cm, rightMargin=2 * cm, topMargin=2 * cm, bottomMargin=2 * cm)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="ReportH1", parent=styles["Heading1"], fontSize=18, spaceAfter=10))
    styles.add(ParagraphStyle(name="ReportH2", parent=styles["Heading2"], fontSize=13, spaceBefore=12, spaceAfter=6))
    styles.add(ParagraphStyle(name="ReportBody", parent=styles["Normal"], fontSize=11, leading=15))

    story = [Paragraph(title.replace("&", "&amp;"), styles["ReportH1"]), Spacer(1, 8)]
    parser = _SimpleHTMLCollector()
    parser.feed(body_html)

    for kind, value in parser.blocks:
        if kind == "img" and value.startswith("data:image"):
            try:
                import base64

                header, b64 = value.split(",", 1)
                raw = base64.b64decode(b64)
                img_buf = io.BytesIO(raw)
                story.append(RLImage(img_buf, width=16 * cm, height=6.2 * cm))
                story.append(Spacer(1, 8))
            except Exception:
                continue
        else:
            safe = (
                value.replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
            )
            style = styles["ReportBody"]
            if kind == "h1":
                style = styles["ReportH1"]
            elif kind in ("h2", "h3"):
                style = styles["ReportH2"]
            elif kind == "li":
                safe = "• " + safe
            story.append(Paragraph(safe, style))

    doc.build(story)
    return buf.getvalue()


def export_pdf_bytes(body_html: str, title: str) -> bytes:
    try:
        return _export_pdf_weasyprint(body_html, title)
    except Exception:
        return _export_pdf_reportlab(body_html, title)


def export_docx_bytes(body_html: str, title: str) -> bytes:
    doc = Document()
    doc.add_heading(title, level=0)

    cleaned = re.sub(r'<img[^>]+src="data:image[^"]+"[^>]*>', "[діаграма]", body_html, flags=re.I)
    plain = _strip_tags(cleaned)
    for block in plain.split("\n"):
        line = block.strip()
        if not line:
            continue
        p = doc.add_paragraph(line)
        for run in p.runs:
            run.font.size = Pt(11)

    for match in re.finditer(r'src="(data:image/[^;]+;base64,([^"]+))"', body_html):
        try:
            import base64

            raw = base64.b64decode(match.group(2))
            stream = io.BytesIO(raw)
            doc.add_picture(stream, width=Inches(5.5))
        except Exception:
            continue

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_report(body_html: str, title: str, fmt: ExportFormat) -> tuple[bytes, str, str]:
    safe_title = re.sub(r"[^\w\-]+", "_", title, flags=re.U)[:60] or "shift_report"
    if fmt == "pdf":
        return export_pdf_bytes(body_html, title), "application/pdf", f"{safe_title}.pdf"
    if fmt == "docx":
        return (
            export_docx_bytes(body_html, title),
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            f"{safe_title}.docx",
        )
    return export_html_bytes(body_html, title), "text/html; charset=utf-8", f"{safe_title}.html"
